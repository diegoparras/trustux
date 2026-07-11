#!/usr/bin/env python3
"""
gen_unlock_fixtures.py — Genera archivos de ejemplo PROTEGIDOS para probar el módulo goodunluck.
Todos son de juguete, sin datos reales. Las contraseñas son conocidas (para los tests).

  unlock/excel-protegido.xlsx   hoja + libro protegidos (Tier 1: quitar restricción)
  unlock/word-restringido.docx  restrict-editing / documentProtection (Tier 1)
  unlock/pdf-permisos.pdf        owner-password: abre sin clave pero con permisos bloqueados (Tier 1)
  unlock/pdf-clave.pdf           user-password "abrime": no abre sin la clave (Tier 2)
  unlock/excel-cifrado.xlsx      cifrado con clave "secreto" (Tier 2)  [requiere msoffcrypto para generar]

Uso:  python scripts/gen_unlock_fixtures.py
Deps: pip install openpyxl python-docx pikepdf
"""
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "fixtures" / "unlock"
OUT.mkdir(parents=True, exist_ok=True)

# --- Excel: hoja + libro protegidos (con clave conocida, pero Tier 1 la quita sin usarla) ---
from openpyxl import Workbook
from openpyxl.workbook.protection import WorkbookProtection
wb = Workbook()
ws = wb.active
ws.title = "Balance"
ws["A1"] = "Concepto"; ws["B1"] = "Importe"
ws["A2"] = "Total activo"; ws["B2"] = 1000000
ws.protection.sheet = True
ws.protection.password = "secreto"          # protección de hoja (hash en <sheetProtection>)
wb.security = WorkbookProtection(workbookPassword="secreto", lockStructure=True)  # protección de libro
wb.save(OUT / "excel-protegido.xlsx")
print("  excel-protegido.xlsx")

# --- Word: restricción de edición (documentProtection en word/settings.xml) ---
from docx import Document
from docx.oxml.ns import qn
doc = Document()
doc.add_heading("Informe confidencial", level=1)
doc.add_paragraph("Importe total: 1.000.000. Documento de prueba de goodunluck.")
settings = doc.settings.element
dp = settings.makeelement(qn("w:documentProtection"), {
    qn("w:edit"): "readOnly", qn("w:enforcement"): "1",
})
settings.append(dp)
doc.save(OUT / "word-restringido.docx")
print("  word-restringido.docx")

# --- PDF: owner-password (permisos bloqueados) y user-password (clave de apertura) ---
import pikepdf
def pdf_base():
    pdf = pikepdf.Pdf.new()
    pdf.add_blank_page(page_size=(612, 792))
    return pdf

perms = pikepdf.Permissions(extract=False, modify_other=False, modify_annotation=False, print_lowres=False, print_highres=False)
pdf_base().save(OUT / "pdf-permisos.pdf",
                encryption=pikepdf.Encryption(owner="ownerpass", user="", allow=perms))
print("  pdf-permisos.pdf  (owner-password: abre sin clave, permisos bloqueados)")

pdf_base().save(OUT / "pdf-clave.pdf",
                encryption=pikepdf.Encryption(owner="ownerpass", user="abrime"))
print("  pdf-clave.pdf  (user-password 'abrime': Tier 2)")

# --- Excel CIFRADO con clave (Tier 2): cifrado agile ECMA-376, clave "secreto" ---
try:
    import io
    from msoffcrypto.format.ooxml import OOXMLFile
    plain = io.BytesIO()
    wb2 = Workbook(); ws2 = wb2.active
    ws2["A1"] = "Concepto"; ws2["B1"] = "Importe"; ws2["A2"] = "Total activo"; ws2["B2"] = 1000000
    wb2.save(plain); plain.seek(0)
    enc = io.BytesIO(); OOXMLFile(plain).encrypt("secreto", enc)
    (OUT / "excel-cifrado.xlsx").write_bytes(enc.getvalue())
    print("  excel-cifrado.xlsx  (cifrado agile, clave 'secreto', Tier 2)")
except Exception as e:
    print(f"  excel-cifrado.xlsx  OMITIDO ({type(e).__name__}: {e})")

# --- ZIP cifrado (AES): el índice (nombres) queda visible sin clave = fuga de metadatos ---
try:
    import pyzipper
    with pyzipper.AESZipFile(OUT / "zip-cifrado.zip", "w", compression=pyzipper.ZIP_DEFLATED, encryption=pyzipper.WZ_AES) as z:
        z.setpassword(b"secreto")
        z.writestr("balance-secreto.txt", b"Total activo: 1000000")
        z.writestr("nomina/sueldos.csv", b"empleado,sueldo\nperez,500000")
    print("  zip-cifrado.zip  (AES clave 'secreto'; nombres visibles sin clave)")
except Exception as e:
    print(f"  zip-cifrado.zip  OMITIDO ({type(e).__name__})")

print("Listo.")
